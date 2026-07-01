from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.document_loader import (
    DocumentLoader,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)


def test_load_txt_preserves_metadata(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Xin chào sinh viên.\nĐây là tài liệu mẫu.", encoding="utf-8")

    documents = DocumentLoader().load(str(file_path), "sample.txt", "file-1")

    assert documents == [
        {
            "text": "Xin chào sinh viên.\nĐây là tài liệu mẫu.",
            "metadata": {
                "file_id": "file-1",
                "file_name": "sample.txt",
                "page": 1,
            },
        }
    ]


def test_load_docx_returns_non_empty_paragraphs_with_metadata(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    docx = DocxDocument()
    docx.add_paragraph("Điều 1. Quy định chung.")
    docx.add_paragraph("")
    docx.add_paragraph("Điều 2. Điều kiện thực tập.")
    docx.save(file_path)

    documents = DocumentLoader().load(str(file_path), "sample.docx", "file-2")

    assert [document["text"] for document in documents] == [
        "Điều 1. Quy định chung.",
        "Điều 2. Điều kiện thực tập.",
    ]
    assert documents[0]["metadata"] == {
        "file_id": "file-2",
        "file_name": "sample.docx",
        "page": None,
        "paragraph_index": 1,
    }
    assert documents[1]["metadata"]["paragraph_index"] == 3


def test_unsupported_file_type_raises_clear_error(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.md"
    file_path.write_text("# Unsupported", encoding="utf-8")

    with pytest.raises(UnsupportedFileTypeError, match="Supported file types"):
        DocumentLoader().load(str(file_path), "sample.md", "file-3")


def test_empty_txt_raises_clear_error(tmp_path: Path) -> None:
    file_path = tmp_path / "empty.txt"
    file_path.write_text("   ", encoding="utf-8")

    with pytest.raises(EmptyDocumentError, match="No extractable text"):
        DocumentLoader().load(str(file_path), "empty.txt", "file-4")
